import re
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import datetime
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager
from time import sleep
import os,sys
import subprocess
try:
    from comtypes import client
except ImportError:
    client = None
#from win32com import client

def resource_path(relative_path):
	if hasattr(sys, "_MEIPASS"):
		base_path = sys._MEIPASS
	else:
		base_path = os.path.abspath(".")
	return os.path.join(base_path, relative_path)

url = "http://172.25.158.199/engineroomcheck/"
resp = requests.get(url)
resp.encoding = 'utf-8' 
print("encoding: %s" % resp.encoding)

soup = BeautifulSoup(resp.text, 'html.parser')
"""for br in soup.find_all("br"):
    br.replace_with("\n")"""
host_name=soup.select("div.ui-block-a")
contents=soup.select("div.ui-block-c")
doc = Document(resource_path('機房工作日誌.docx'))
print(resource_path('機房工作日誌.docx成功讀取'))
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')

before=["C:","D:","E:","F:","G:"]
after=["C槽(GB):","D槽(GB):","E槽(GB):","F槽(GB):","G槽(GB):"]
for i in range(1, 21):
	print (host_name[i].text,contents[i].text)
	s=""
	string=contents[i].get_text()
	regex = re.compile(r'[A-Z]:\w{,6}')
	matchs = regex.findall(string)
	for index,match in enumerate(matchs):
		for n in range(0,5):
			match=match.replace(before[n],after[n])
		s=s+match
		#print("index:"+str(index)+"len:"+str(len(matchs)))
		if(index<len(matchs)-1):
			s=s+"\n"
	print(s)
	cell = doc.tables[1].cell(i, 4)
	cell.text = s	
	cell_font = cell.paragraphs[0].runs[0].font
	cell_font.name = "標楷體"
	cell_font.size = Pt(8)	
	print ("\n")

chrome_options = Options()
chrome_options.add_argument('--headless')
chrome_options.add_argument('--no-sandbox')
chrome_options.add_argument('--disable-gpu')
chrome_options.add_argument('--disable-dev-shm-usage')

service=Service(ChromeDriverManager().install())
driver = webdriver.Chrome(options=chrome_options,service=service)
#driver = webdriver.Ie("IEDriverServer.exe")
#driver = webdriver.Chrome()
#driver.implicitly_wait(10)
driver.get('http://191.7.1.31/')
sleep(5)
soup = BeautifulSoup(driver.page_source, 'html.parser')
content1=soup.select("div.ch1")
content2=soup.select("div.ch2")
print("溫度:"+content1[0].text)
print("濕度:"+content2[0].text)

cell=doc.tables[1].cell(35, 4)
cell.text="溫度:"+content1[0].text+" 濕度:"+content2[0].text
cell_font = cell.paragraphs[0].runs[0].font
cell_font.name = "標楷體"
cell_font.size = Pt(7)
driver.close()
driver.quit()

d = datetime.datetime.now()

doc.save("./"+str(d.year-1911)+"{:>02d}".format(d.month)+"{:>02d}".format(d.day)+"-機房工作日誌.docx")


def convert_doc_to_pdf(in_file, out_file):
    wd_format_pdf = 17
    word = client.CreateObject('Word.Application')
    doc = word.Documents.Open(in_file)
    doc.SaveAs(out_file, FileFormat=wd_format_pdf)
    doc.Close()
    word.Quit()

def doc2pdf(doc_name, pdf_name):
    try:
        word = client.DispatchEx("Word.Application")
        worddoc = word.Documents.Open(doc_name,ReadOnly = 1)
        worddoc.SaveAs(pdf_name, FileFormat = 17)
        worddoc.Close()
        return pdf_name
    except:
        return 1
def doc2pdf_linux(doc):
    cmd = 'libreoffice --infilter="Microsoft Word 2007/2010/2013 XML" --convert-to pdf:writer_pdf_Export '.split() + [doc]
    p = subprocess.Popen(cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE)
    p.wait()
    stdout, stderr = p.communicate()
    if stderr:
        raise subprocess.SubprocessError(stderr)

in_file = os.getcwd()+"/"+str(d.year-1911)+"{:>02d}".format(d.month)+"{:>02d}".format(d.day)+"-機房工作日誌.docx"
out_file = os.getcwd()+"/"+str(d.year-1911)+"{:>02d}".format(d.month)+"{:>02d}".format(d.day)+"-機房工作日誌.pdf"
print(in_file,out_file)
if client is None:
	doc2pdf_linux(in_file)
else:
	convert_doc_to_pdf(in_file, out_file)
print(in_file+"已經轉換"+out_file)

os.remove(str(d.year-1911)+"{:>02d}".format(d.month)+"{:>02d}".format(d.day)+"-機房工作日誌.docx")
print(in_file+"已刪除")
